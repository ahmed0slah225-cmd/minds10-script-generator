# -*- coding: utf-8 -*-
"""
Minds10 Script Generator — خط إنتاج كامل
==========================================
بدل ما يكون فيه زرار واحد "اكتب السكريبت"، التطبيق ده بقى خط إنتاج على
مراحل، كل مرحلة بتاخد وقتها وتقدر تراجعها قبل ما تكمل:

الموضوع + المصادر → استخراج المعرفة → بناء الهيكل → توليد الهوكس
→ كتابة كل جزء لوحده (مع مراجعة احتفاظ) → مراجعات نهائية (هوك/احتفاظ/
حقائق/أسلوب) → تحرير نهائي + B-roll وتعليمات مونتاج.

كل التعليمات الخاصة بشخصية القناة موجودة في SCRIPT_RULES.md (جذر
المشروع) - عدّل هناك عشان تغيّر الأسلوب، من غير ما تلمس الكود.
كل منطق الـ Prompts وخط الإنتاج نفسه موجود في engine/.
"""

import io
from datetime import datetime

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from engine.knowledge_base import KnowledgeBase, SOURCE_TYPES
from engine.rules import get_script_rules, save_script_rules
from engine.gemini_client import make_client, GenerationError
from engine import pipeline

# ----------------------------- إعدادات عامة ----------------------------- #

st.set_page_config(page_title="Minds10 - خط إنتاج السكريبتات", page_icon="🎬", layout="wide")

MODEL_OPTIONS = [
    "gemini-3.7-flash",
    "gemini-3.6-flash",
    "gemini-3.5-flash",
    "gemini-2.5-flash",
]

STAGE_NAMES = {
    1: "١. الموضوع والمصادر",
    2: "٢. استخراج المعرفة",
    3: "٣. الهيكل",
    4: "٤. الهوكس",
    5: "٥. كتابة السكريبت",
    6: "٦. المراجعة",
    7: "٧. النهائي",
}

_DEFAULTS = {
    "stage": 1,
    "kb": None,
    "topic": "",
    "duration_min": 30,
    "audience": "",
    "tone": "تشويقي / إثارة فضول",
    "extraction": None,
    "outline": None,
    "hooks_data": None,
    "chosen_hook_id": None,
    "sections_text": {},
    "retention_reviews": {},
    "final_reviews": {},
    "final_script": None,
    "broll": None,
}
for k, v in _DEFAULTS.items():
    st.session_state.setdefault(k, v)
if st.session_state["kb"] is None:
    st.session_state["kb"] = KnowledgeBase()


# ----------------------------- دوال مساعدة عامة ----------------------------- #

def get_secret_api_key() -> str:
    try:
        return st.secrets.get("GEMINI_API_KEY", "")
    except Exception:
        return ""


def get_api_key() -> str:
    return get_secret_api_key() or st.session_state.get("manual_api_key", "")


def go_to(stage: int):
    st.session_state["stage"] = stage


def set_paragraph_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def build_docx(script_text: str, topic: str, extra_sections: dict = None) -> io.BytesIO:
    doc = Document()
    title = doc.add_heading(f"سكريبت: {topic}", level=0)
    set_paragraph_rtl(title)

    for raw_line in script_text.split("\n"):
        line = raw_line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("## "):
            p = doc.add_heading(line.replace("## ", "").strip(), level=2)
        elif line.startswith("# "):
            p = doc.add_heading(line.replace("# ", "").strip(), level=1)
        else:
            p = doc.add_paragraph(line)
        set_paragraph_rtl(p)

    if extra_sections:
        for heading, text in extra_sections.items():
            h = doc.add_heading(heading, level=1)
            set_paragraph_rtl(h)
            for line in text.split("\n"):
                p = doc.add_paragraph(line)
                set_paragraph_rtl(p)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def word_count(text: str) -> int:
    import re
    return len(re.findall(r"\S+", text))


def run_safely(fn, *args, **kwargs):
    """بينفذ نداء API ويرجع (result, error_message)."""
    try:
        with st.spinner("بيشتغل... ⏳"):
            return fn(*args, **kwargs), None
    except GenerationError as e:
        return None, str(e)
    except Exception as e:
        return None, f"حصل خطأ غير متوقع: {e}"


# ----------------------------- الشريط الجانبي ----------------------------- #

with st.sidebar:
    st.header("⚙️ الإعدادات")

    if not get_secret_api_key():
        st.text_input("مفتاح Gemini API", type="password", key="manual_api_key")
    else:
        st.success("مفتاح الـ API متظبط من إعدادات السيرفر ✅")

    model = st.selectbox("الموديل", options=MODEL_OPTIONS, index=0)

    st.session_state["duration_min"] = st.slider(
        "مدة الفيديو (دقايق)", min_value=10, max_value=45,
        value=st.session_state["duration_min"], step=5,
    )
    st.caption(f"هيتكتب تقريبًا {st.session_state['duration_min'] * pipeline.prompts.WORDS_PER_MINUTE:,} كلمة")

    st.divider()
    with st.expander("📖 قواعد أسلوب القناة (SCRIPT_RULES.md)"):
        st.caption("عدّل هنا وشخصية القناة هتتغيّر في كل السكريبتات الجديدة.")
        rules_editable = st.text_area("محتوى الملف", value=get_script_rules(), height=300, key="rules_editor")
        if st.button("💾 احفظ القواعد"):
            if save_script_rules(rules_editable):
                st.success("اتحفظ ✅")
            else:
                st.error("تعذّر الحفظ.")

    st.divider()
    st.subheader("مراحل الإنتاج")
    for k, name in STAGE_NAMES.items():
        if k == st.session_state["stage"]:
            st.markdown(f"**🔵 {name}**")
        elif k < st.session_state["stage"]:
            st.markdown(f"✅ {name}")
        else:
            st.markdown(f"⚪ {name}")

    if st.button("🔄 ابدأ مشروع جديد", use_container_width=True):
        for k, v in _DEFAULTS.items():
            st.session_state[k] = v
        st.session_state["kb"] = KnowledgeBase()
        st.rerun()


api_key = get_api_key()
rules_text = get_script_rules()
stage = st.session_state["stage"]

st.title("🎬 Minds10 — خط إنتاج السكريبتات")
st.caption("من الموضوع لحد السكريبت الجاهز، مرحلة مرحلة، مبني على مصادرك انت بس.")
st.divider()


def require_api_key() -> bool:
    if not api_key:
        st.error("محتاج تحط مفتاح Gemini API الأول (في الشريط الجانبي).")
        return False
    return True


# =============================== المرحلة ١: الموضوع والمصادر =============================== #

if stage == 1:
    st.subheader("١. الموضوع، الإعدادات، وقاعدة المعرفة")

    col1, col2 = st.columns([2, 1])
    with col1:
        st.session_state["topic"] = st.text_area(
            "📌 موضوع الفيديو", value=st.session_state["topic"],
            placeholder="مثلاً: ليه بعض الناس بتفقد شغفها بسرعة؟", height=90,
        )
    with col2:
        st.session_state["audience"] = st.text_input("👥 الجمهور المستهدف (اختياري)", value=st.session_state["audience"])
        st.session_state["tone"] = st.selectbox(
            "🎭 نبرة الفيديو",
            options=["تشويقي / إثارة فضول", "تعليمي مبسّط", "قصصي حكواتي", "تحليلي عميق"],
            index=["تشويقي / إثارة فضول", "تعليمي مبسّط", "قصصي حكواتي", "تحليلي عميق"].index(st.session_state["tone"]),
        )

    st.divider()
    st.markdown("### 🗂️ قاعدة المعرفة (المصادر)")
    st.caption(
        "البرنامج مش هيخترع معلومات ولا يدوّر بنفسه على الإنترنت - "
        "هو هيعتمد على المصادر اللي انت هتضيفها هنا بس (كتب، مقالات، أبحاث، "
        "نصوص فيديوهات مرجعية، ملاحظاتك الخاصة)."
    )

    up_col, type_col = st.columns([3, 1])
    with up_col:
        uploaded_files = st.file_uploader(
            "ارفع ملفات (TXT / MD / PDF)", type=["txt", "md", "pdf"], accept_multiple_files=True,
        )
    with type_col:
        upload_type = st.selectbox("نوع المصادر المرفوعة", options=SOURCE_TYPES, index=1)

    if uploaded_files and st.button("➕ أضف الملفات المرفوعة لقاعدة المعرفة"):
        added = 0
        for f in uploaded_files:
            src = st.session_state["kb"].add_uploaded_file(f, source_type=upload_type)
            if src:
                added += 1
        st.success(f"اتضاف {added} مصدر ✅")
        st.rerun()

    with st.expander("✏️ أضف ملاحظة أو نص يدويًا"):
        note_title = st.text_input("عنوان المصدر", key="note_title")
        note_type = st.selectbox("نوعه", options=SOURCE_TYPES, index=4, key="note_type")
        note_content = st.text_area("المحتوى", height=150, key="note_content")
        if st.button("➕ أضف كمصدر"):
            if note_content.strip():
                st.session_state["kb"].add_text(note_title or "ملاحظة", note_content, note_type)
                st.success("اتضاف ✅")
                st.rerun()
            else:
                st.warning("اكتب محتوى الملاحظة الأول.")

    sources = st.session_state["kb"].summary_table()
    if sources:
        st.markdown(f"**المصادر الحالية ({len(sources)}):**")
        for s in sources:
            c1, c2, c3 = st.columns([4, 2, 1])
            c1.write(f"📄 {s['title']}")
            c2.caption(f"{s['type']} — {s['chars']:,} حرف")
            if c3.button("🗑️", key=f"del_{s['id']}"):
                st.session_state["kb"].remove(s["id"])
                st.rerun()
    else:
        st.info("لسه معندكش مصادر. تقدر تكمل من غيرها بس النتيجة هتكون أعم وأقل دقة في الأرقام والدراسات.")

    st.divider()
    disabled = not st.session_state["topic"].strip()
    if st.button("➡️ التالي: استخراج المعرفة", type="primary", disabled=disabled, use_container_width=True):
        go_to(2)
        st.rerun()


# =============================== المرحلة ٢: استخراج المعرفة =============================== #

elif stage == 2:
    st.subheader("٢. استخراج المعرفة من المصادر")
    st.caption("الموديل هيقرا المصادر ويطلع منها الفكرة الأساسية، النقاط، القصص، الأرقام، الاقتباسات، وأي حاجة محتاجة تحقق.")

    if st.button("🔍 استخرج المعرفة", type="primary") and require_api_key():
        client = make_client(api_key)
        kb_context = st.session_state["kb"].to_context_text(query=st.session_state["topic"])
        result, err = run_safely(
            pipeline.extract_knowledge, client, model, st.session_state["topic"], kb_context, rules_text
        )
        if err:
            st.error(err)
        else:
            st.session_state["extraction"] = result
            st.success("تم الاستخراج ✅")

    extraction = st.session_state["extraction"]
    if extraction:
        st.markdown(f"**💡 الفكرة الأساسية:** {extraction.get('main_idea', '')}")

        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**🔑 أهم النقاط:**")
            for p in extraction.get("key_points", []):
                st.markdown(f"- {p}")
            st.markdown("**📚 القصص:**")
            for s in extraction.get("stories", []):
                st.markdown(f"- **{s.get('title', '')}**: {s.get('summary', '')} _(المصدر: {s.get('source', '؟')})_")
        with c2:
            st.markdown("**📊 الأرقام والدراسات:**")
            for s in extraction.get("stats", []):
                st.markdown(f"- {s.get('stat', '')} _(المصدر: {s.get('source', '؟')})_")
            st.markdown("**💬 الاقتباسات:**")
            for q in extraction.get("quotes", []):
                st.markdown(f"- \"{q.get('quote', '')}\" — {q.get('attributed_to', '؟')} _(المصدر: {q.get('source', '؟')})_")

        needs_verif = extraction.get("needs_verification", [])
        if needs_verif:
            with st.expander("⚠️ معلومات محتاجة تحقق يدوي"):
                for n in needs_verif:
                    st.markdown(f"- {n}")

    st.divider()
    b1, b2 = st.columns(2)
    if b1.button("⬅️ رجوع", use_container_width=True):
        go_to(1)
        st.rerun()
    if b2.button("➡️ التالي: بناء الهيكل", type="primary", disabled=not extraction, use_container_width=True):
        go_to(3)
        st.rerun()


# =============================== المرحلة ٣: الهيكل =============================== #

elif stage == 3:
    st.subheader("٣. بناء هيكل السكريبت (Outline)")
    st.caption("Hook → مشكلة → فضول → قصة → تفسير → مثال → مفاجأة → حل → تطبيق → نهاية")

    if st.button("🏗️ ابنِ الهيكل", type="primary") and require_api_key():
        client = make_client(api_key)
        result, err = run_safely(
            pipeline.build_outline, client, model, st.session_state["topic"],
            st.session_state["duration_min"], rules_text, st.session_state["extraction"],
        )
        if err:
            st.error(err)
        else:
            st.session_state["outline"] = result
            st.success("تم بناء الهيكل ✅")

    outline = st.session_state["outline"]
    if outline:
        for section in outline:
            with st.container(border=True):
                c1, c2 = st.columns([3, 1])
                c1.markdown(f"**{section.get('title', section.get('key'))}**")
                c1.caption(section.get("goal", ""))
                if section.get("notes"):
                    c1.caption(f"📝 {section.get('notes')}")
                c2.metric("كلمات", section.get("target_words", "-"))

    st.divider()
    b1, b2 = st.columns(2)
    if b1.button("⬅️ رجوع", use_container_width=True):
        go_to(2)
        st.rerun()
    if b2.button("➡️ التالي: توليد الهوكس", type="primary", disabled=not outline, use_container_width=True):
        go_to(4)
        st.rerun()


# =============================== المرحلة ٤: الهوكس =============================== #

elif stage == 4:
    st.subheader("٤. توليد وتقييم الهوكس")
    st.caption("هيتكتب 10 نسخ من الهوك ويتقيّموا على: الفضول، قوة أول جملة، فتح Loop، عدم كشف الإجابة، الاحتفاظ.")

    if st.button("✨ ولّد 10 هوكس", type="primary") and require_api_key():
        client = make_client(api_key)
        result, err = run_safely(
            pipeline.generate_hooks, client, model, st.session_state["topic"],
            st.session_state["extraction"], rules_text, 10,
        )
        if err:
            st.error(err)
        else:
            st.session_state["hooks_data"] = result
            st.session_state["chosen_hook_id"] = result.get("recommended_id")
            st.success("تم توليد الهوكس ✅")

    hooks_data = st.session_state["hooks_data"]
    if hooks_data:
        hooks = sorted(hooks_data.get("hooks", []), key=lambda h: h.get("total", 0), reverse=True)
        recommended_id = hooks_data.get("recommended_id")

        options = {h["id"]: h for h in hooks}
        chosen_id = st.radio(
            "اختار الهوك اللي هتستخدمه (مرتّبين من الأعلى تقييمًا):",
            options=list(options.keys()),
            index=list(options.keys()).index(st.session_state["chosen_hook_id"])
            if st.session_state["chosen_hook_id"] in options else 0,
            format_func=lambda hid: f"⭐ #{hid} — {options[hid]['total']}/50 نقطة" + (" 🏆 مُوصى به" if hid == recommended_id else ""),
        )
        st.session_state["chosen_hook_id"] = chosen_id

        chosen = options[chosen_id]
        st.markdown(f"> {chosen['text']}")
        sc = chosen.get("scores", {})
        cols = st.columns(5)
        labels = {"curiosity": "فضول", "first_line_strength": "قوة أول جملة",
                  "open_loop": "فتح Loop", "no_answer_reveal": "عدم كشف الإجابة",
                  "retention_potential": "احتفاظ"}
        for i, (key, label) in enumerate(labels.items()):
            cols[i].metric(label, f"{sc.get(key, '-')}/10")

        with st.expander("شوف باقي الهوكس"):
            for h in hooks:
                st.markdown(f"**#{h['id']} ({h['total']}/50):** {h['text']}")

    st.divider()
    b1, b2 = st.columns(2)
    if b1.button("⬅️ رجوع", use_container_width=True):
        go_to(3)
        st.rerun()
    if b2.button("➡️ التالي: كتابة السكريبت", type="primary", disabled=not hooks_data, use_container_width=True):
        go_to(5)
        st.rerun()


# =============================== المرحلة ٥: كتابة السكريبت جزء جزء =============================== #

elif stage == 5:
    st.subheader("٥. كتابة السكريبت — جزء جزء")
    st.caption("كل جزء بيتكتب لوحده، وبعد كل جزء بتقدر تشغّل مراجع الاحتفاظ يشوف فين ممكن المشاهد يزهق.")

    outline = st.session_state["outline"]
    hooks_data = st.session_state["hooks_data"]
    chosen_hook_text = next(
        (h["text"] for h in hooks_data.get("hooks", []) if h["id"] == st.session_state["chosen_hook_id"]), ""
    )

    writable_sections = [s for s in outline if s.get("key") != "hook"]
    done, total = pipeline.outline_progress(outline, st.session_state["sections_text"])
    st.progress(done / total if total else 0, text=f"{done} من {total} أجزاء اتكتبت")

    st.markdown("### 🎣 الهوك المُختار")
    st.info(chosen_hook_text)

    client = make_client(api_key) if api_key else None
    prev_texts_accum = [chosen_hook_text]

    for section in writable_sections:
        key = section["key"]
        with st.expander(f"### {section.get('title')}", expanded=key not in st.session_state["sections_text"]):
            st.caption(section.get("goal", ""))
            current_text = st.session_state["sections_text"].get(key, "")

            c1, c2 = st.columns([1, 1])
            if c1.button(f"🎬 {'أعد كتابة' if current_text else 'اكتب'} هذا الجزء", key=f"write_{key}"):
                if require_api_key():
                    prev_text = "\n\n".join(prev_texts_accum)
                    result, err = run_safely(
                        pipeline.write_section, client, model, section, st.session_state["topic"],
                        rules_text, st.session_state["extraction"], prev_text, chosen_hook_text,
                    )
                    if err:
                        st.error(err)
                    else:
                        st.session_state["sections_text"][key] = result
                        st.session_state["retention_reviews"].pop(key, None)
                        st.rerun()

            if current_text and c2.button("🔍 راجع الاحتفاظ", key=f"review_{key}"):
                if require_api_key():
                    result, err = run_safely(pipeline.review_retention, client, model, current_text, rules_text)
                    if err:
                        st.error(err)
                    else:
                        st.session_state["retention_reviews"][key] = result

            edited = st.text_area("النص (تقدر تعدّله يدويًا)", value=current_text, height=200, key=f"text_{key}")
            if edited != current_text:
                st.session_state["sections_text"][key] = edited

            review = st.session_state["retention_reviews"].get(key)
            if review:
                risk = review.get("overall_risk", "؟")
                risk_emoji = {"منخفض": "🟢", "متوسط": "🟡", "مرتفع": "🔴"}.get(risk, "⚪")
                st.markdown(f"**تقييم الاحتفاظ:** {risk_emoji} {risk}")
                for issue in review.get("issues", []):
                    st.warning(f"**المشكلة:** {issue.get('problem')}\n\n**المقطع:** {issue.get('excerpt')}\n\n**الاقتراح:** {issue.get('fix')}")

        if current_text:
            prev_texts_accum.append(current_text)

    st.divider()
    b1, b2 = st.columns(2)
    if b1.button("⬅️ رجوع", use_container_width=True):
        go_to(4)
        st.rerun()
    if b2.button("➡️ التالي: المراجعة النهائية", type="primary", disabled=done < total, use_container_width=True):
        go_to(6)
        st.rerun()


# =============================== المرحلة ٦: المراجعة النهائية =============================== #

elif stage == 6:
    st.subheader("٦. فريق المراجعين")
    st.caption("4 مراجعين متخصصين بيشتغلوا على السكريبت كامل قبل التحرير النهائي.")

    hooks_data = st.session_state["hooks_data"]
    chosen_hook_text = next(
        (h["text"] for h in hooks_data.get("hooks", []) if h["id"] == st.session_state["chosen_hook_id"]), ""
    )
    full_script = pipeline.assemble_script(chosen_hook_text, st.session_state["sections_text"], st.session_state["outline"])
    st.session_state["_assembled_script"] = full_script

    with st.expander("📄 شوف السكريبت الكامل (قبل التحرير النهائي)"):
        st.markdown(full_script)

    client = make_client(api_key) if api_key else None
    reviews = st.session_state["final_reviews"]

    rc1, rc2, rc3, rc4 = st.columns(4)
    if rc1.button("🎣 مراجعة الهوك", use_container_width=True) and require_api_key():
        result, err = run_safely(pipeline.review_hook, client, model, chosen_hook_text, rules_text)
        if err:
            st.error(err)
        else:
            reviews["hook"] = result
    if rc2.button("📉 مراجعة الاحتفاظ (كامل)", use_container_width=True) and require_api_key():
        result, err = run_safely(pipeline.review_retention, client, model, full_script, rules_text)
        if err:
            st.error(err)
        else:
            reviews["retention"] = result
    if rc3.button("🔎 تدقيق الحقائق", use_container_width=True) and require_api_key():
        kb_context = st.session_state["kb"].to_context_text(query=st.session_state["topic"])
        result, err = run_safely(pipeline.fact_check, client, model, full_script, kb_context, st.session_state["extraction"])
        if err:
            st.error(err)
        else:
            reviews["fact_check"] = result
    if rc4.button("🎨 مراجعة الأسلوب", use_container_width=True) and require_api_key():
        result, err = run_safely(pipeline.review_style, client, model, full_script, rules_text)
        if err:
            st.error(err)
        else:
            reviews["style"] = result

    if reviews.get("hook"):
        with st.expander("🎣 نتيجة مراجعة الهوك", expanded=True):
            st.json(reviews["hook"])
    if reviews.get("retention"):
        with st.expander("📉 نتيجة مراجعة الاحتفاظ"):
            st.json(reviews["retention"])
    if reviews.get("fact_check"):
        with st.expander("🔎 نتيجة تدقيق الحقائق"):
            st.json(reviews["fact_check"])
    if reviews.get("style"):
        with st.expander("🎨 نتيجة مراجعة الأسلوب"):
            st.json(reviews["style"])

    st.divider()
    st.markdown("### 🖋️ التحرير النهائي")
    st.caption("المحرر هياخد كل الملاحظات اللي جمعتها فوق وينتج نسخة نهائية واحدة.")
    if st.button("🖋️ ولّد النسخة النهائية", type="primary") and require_api_key():
        result, err = run_safely(pipeline.final_edit, client, model, full_script, reviews, rules_text)
        if err:
            st.error(err)
        else:
            st.session_state["final_script"] = result
            st.success("تم ✅")

    st.divider()
    b1, b2 = st.columns(2)
    if b1.button("⬅️ رجوع", use_container_width=True):
        go_to(5)
        st.rerun()
    if b2.button("➡️ التالي: السكريبت النهائي", type="primary",
                 disabled=not st.session_state["final_script"], use_container_width=True):
        go_to(7)
        st.rerun()


# =============================== المرحلة ٧: النهائي =============================== #

elif stage == 7:
    st.subheader("٧. السكريبت النهائي")

    final_script = st.session_state["final_script"]
    topic = st.session_state["topic"]
    wc = word_count(final_script)
    est_minutes = round(wc / pipeline.prompts.WORDS_PER_MINUTE, 1)
    st.caption(f"عدد الكلمات: {wc:,} — المدة التقريبية: {est_minutes} دقيقة")
    st.markdown(final_script)

    st.divider()
    client = make_client(api_key) if api_key else None
    if st.button("🎬 ولّد اقتراحات B-roll وتعليمات المونتاج") and require_api_key():
        result, err = run_safely(pipeline.suggest_broll, client, model, final_script)
        if err:
            st.error(err)
        else:
            st.session_state["broll"] = result

    broll_text_for_docx = ""
    if st.session_state["broll"]:
        st.markdown("### 🎥 اقتراحات B-roll وتعليمات المونتاج")
        lines = []
        for sec in st.session_state["broll"].get("sections", []):
            with st.container(border=True):
                st.markdown(f"**{sec.get('section')}**")
                for idea in sec.get("broll_ideas", []):
                    st.markdown(f"- {idea}")
                st.caption(f"🎞️ {sec.get('editing_notes', '')}")
            lines.append(f"{sec.get('section')}:\n" + "\n".join(f"- {i}" for i in sec.get("broll_ideas", [])) + f"\nمونتاج: {sec.get('editing_notes', '')}")
        broll_text_for_docx = "\n\n".join(lines)

    fact_check_data = st.session_state["final_reviews"].get("fact_check")
    if fact_check_data:
        with st.expander("📚 مصادر المعلومات في السكريبت"):
            for c in fact_check_data.get("verified_claims", []):
                st.markdown(f"✅ {c.get('claim')} — _{c.get('source')}_")
            for c in fact_check_data.get("unverified_claims", []):
                st.markdown(f"⚠️ غير موثّق: {c}")

    st.divider()
    dcol1, dcol2 = st.columns(2)
    with dcol1:
        st.download_button(
            "⬇️ تحميل كـ TXT", data=final_script.encode("utf-8"),
            file_name=f"{topic[:40] or 'script'}.txt", mime="text/plain", use_container_width=True,
        )
    with dcol2:
        extra = {"B-roll وتعليمات المونتاج": broll_text_for_docx} if broll_text_for_docx else None
        docx_buffer = build_docx(final_script, topic or "سكريبت", extra_sections=extra)
        st.download_button(
            "⬇️ تحميل كـ Word (docx)", data=docx_buffer,
            file_name=f"{topic[:40] or 'script'}_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    st.divider()
    if st.button("⬅️ رجوع للمراجعة"):
        go_to(6)
        st.rerun()
